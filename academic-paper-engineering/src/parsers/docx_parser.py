"""DOCX 文档解析器 - 将 Word 文档解析为 Document IR"""

import json
from pathlib import Path
from typing import Dict, List, Optional


class DocxParser:
    """解析 .docx 格式的学术文档"""

    def __init__(self):
        self.ir = {
            "document_id": "",
            "metadata": {"source_format": "docx", "source_language": ""},
            "title": {},
            "authors": [],
            "affiliations": [],
            "abstract": {},
            "keywords": {},
            "sections": [],
            "figures": [],
            "tables": [],
            "equations": [],
            "citations": [],
            "references": [],
            "footnotes": []
        }

    def parse(self, file_path: str) -> Dict:
        """解析 DOCX 文件，返回 Document IR"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        self.ir["document_id"] = path.stem

        try:
            from docx import Document
            doc = Document(str(path))

            self._extract_metadata(doc)
            self._extract_title(doc)
            self._extract_authors(doc)
            self._extract_abstract(doc)
            self._extract_sections(doc)
            self._extract_tables(doc)
            self._extract_figures(doc)

        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        return self.ir

    def _extract_metadata(self, doc):
        """提取文档元数据"""
        core_props = doc.core_properties
        self.ir["metadata"]["page_count"] = 0
        if core_props.author:
            self.ir["metadata"]["author"] = core_props.author

    def _extract_title(self, doc):
        """提取标题（通常是第一个 Heading 1 或首段）"""
        for para in doc.paragraphs:
            if para.style and 'Heading' in str(para.style.name) and '1' in str(para.style.name):
                self.ir["title"] = {
                    "id": "title_001",
                    "text": para.text.strip()
                }
                return
        if doc.paragraphs:
            self.ir["title"] = {
                "id": "title_001",
                "text": doc.paragraphs[0].text.strip()
            }

    def _extract_authors(self, doc):
        """提取作者信息"""
        core_props = doc.core_properties
        if core_props.author:
            self.ir["authors"].append({
                "id": "author_001",
                "name": core_props.author,
                "affiliation_id": ""
            })

    def _extract_abstract(self, doc):
        """提取摘要"""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if 'abstract' in text or '摘要' in text:
                abstract_text = para.text.strip()
                if len(abstract_text) < 50 and i + 1 < len(doc.paragraphs):
                    abstract_text = doc.paragraphs[i + 1].text.strip()
                self.ir["abstract"] = {
                    "id": "abstract_001",
                    "text": abstract_text
                }
                return

    def _extract_sections(self, doc):
        """提取章节结构"""
        current_section = None
        section_counter = 0
        para_counter = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = str(para.style.name) if para.style else ""

            if 'Heading 1' in style_name or 'Heading 2' in style_name:
                level = 1 if '1' in style_name else 2
                section_counter += 1
                current_section = {
                    "id": f"section_{section_counter:03d}",
                    "number": str(section_counter),
                    "title": text,
                    "level": level,
                    "paragraphs": [],
                    "subsections": []
                }
                self.ir["sections"].append(current_section)
            elif current_section and text:
                para_counter += 1
                current_section["paragraphs"].append(text)

    def _extract_tables(self, doc):
        """提取表格"""
        for i, table in enumerate(doc.tables):
            rows = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append({
                    "row": row_idx,
                    "cells": [{"text": c} for c in cells]
                })

            self.ir["tables"].append({
                "id": f"table_{i+1:03d}",
                "number": i + 1,
                "caption": "",
                "label": "",
                "headers": [rows[0]] if rows else [],
                "rows": rows[1:] if len(rows) > 1 else []
            })

    def _extract_figures(self, doc):
        """提取图片引用（实际图片文件需通过 docx 解包获取）"""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '图' in text and ('Fig' in text or text.startswith('图')):
                self.ir["figures"].append({
                    "id": f"figure_{i+1:03d}",
                    "number": len(self.ir["figures"]) + 1,
                    "caption": text,
                    "label": ""
                })
